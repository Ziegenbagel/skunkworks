from __future__ import annotations

import os
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Skunkworks_Operator_Manual.docx"
FIGURE = ROOT / "assets" / "mission-control-dashboard-numbered.png"
SETTINGS_FIGURE = ROOT / "assets" / "settings-workspace-numbered.png"
WORKSPACE_ASSETS = ROOT / "assets" / "workspace-diagrams"
SCREEN_ASSETS = ROOT / "assets" / "screenshots"
SYNTHETIC_SCREEN_ASSETS = ROOT / "assets" / "synthetic-screenshots"
MANNY_DIAGRAM = ROOT / "assets" / "manny-warranty-diagram.png"
DEFAULT_SCREENSHOT = SYNTHETIC_SCREEN_ASSETS / "mission-control.png"
SCREENSHOT = Path(os.environ.get("SKUNKWORKS_GUIDE_SCREENSHOT", DEFAULT_SCREENSHOT))
DEFAULT_SETTINGS_SCREENSHOT = SYNTHETIC_SCREEN_ASSETS / "settings-policy.png"
SETTINGS_SCREENSHOT = Path(os.environ.get("SKUNKWORKS_GUIDE_SETTINGS_SCREENSHOT", DEFAULT_SETTINGS_SCREENSHOT))

SCREENSHOTS = {
    key: SYNTHETIC_SCREEN_ASSETS / f"{key}.png"
    for key in (
        "mission-control", "fleet", "galaxy", "navigation-travel",
        "navigation-scan", "resources", "missions", "production", "safety",
        "logbook", "manual-build",
        "manual-field", "manual-cargo", "manual-network",
        "manual-container", "manual-asteroid", "settings-policy",
        "settings-planner", "settings-targets", "settings-floors",
        "settings-status", "settings-roles", "settings-reserve",
        "settings-tanker", "settings-transport",
    )
}

NAVY = RGBColor(5, 23, 34)
CYAN = RGBColor(0, 196, 220)
BLUE = RGBColor(27, 105, 145)
MUTED = RGBColor(91, 111, 123)
AMBER = RGBColor(180, 121, 0)
RED = RGBColor(165, 38, 45)
LIGHT = RGBColor(232, 238, 245)


def font(run, size=11, bold=False, color=NAVY, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True)
        font(p.add_run(text[len(bold_lead):]))
    else:
        font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        font(p.add_run(item))


def add_steps(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "right")):
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        level.append(element)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_pr.append(ilvl)
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_element)
        font(p.add_run(item))


def add_note(doc, title, text, severity="note"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    fill = "FFF4D6" if severity == "warning" else "E8EEF5"
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(title.upper() + "  "), bold=True, color=AMBER if severity == "warning" else BLUE)
    font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def make_screen_crop(key, crop=(0.01, 0.06, 0.99, 0.96)):
    """Normalize supplied UI captures into sharp, manual-ready figures."""

    source = SCREENSHOTS[key]
    if not source.exists():
        return None
    SCREEN_ASSETS.mkdir(parents=True, exist_ok=True)
    image = Image.open(source).convert("RGB")
    width, height = image.size
    left, top, right, bottom = crop
    image = image.crop((
        int(width * left), int(height * top),
        int(width * right), int(height * bottom),
    ))
    # Full-resolution captures are wasteful in DOCX.  This retains more than
    # enough detail for a 6.7-inch print figure while keeping the manual small.
    if image.width > 2100:
        image.thumbnail((2100, 1400), Image.Resampling.LANCZOS)
    output = SCREEN_ASSETS / f"{key}.jpg"
    image.save(output, quality=88, optimize=True)
    return output


def add_screen_figure(doc, key, caption, crop=(0.01, 0.06, 0.99, 0.96), width=6.7):
    path = make_screen_crop(key, crop)
    if path is None:
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = False
    font(paragraph.add_run(caption), size=8.5, color=MUTED)


def make_dashboard_figure():
    if not SCREENSHOT.exists():
        return False
    FIGURE.parent.mkdir(parents=True, exist_ok=True)
    image = Image.open(SCREENSHOT).convert("RGBA")
    draw = ImageDraw.Draw(image, "RGBA")
    try:
        label_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 30)
    except OSError:
        label_font = ImageFont.load_default()
    # Coordinates are fractions so the overlay survives normal window-size changes.
    points = [
        (0.085, 0.165), (0.750, 0.120), (0.515, 0.166), (0.505, 0.425),
        (0.087, 0.280), (0.088, 0.500), (0.087, 0.720), (0.560, 0.775),
        (0.405, 0.895), (0.725, 0.895),
    ]
    radius = max(20, image.width // 70)
    for number, (xf, yf) in enumerate(points, 1):
        x, y = int(image.width * xf), int(image.height * yf)
        draw.ellipse((x-radius, y-radius, x+radius, y+radius), fill=(0, 196, 220, 235), outline=(255, 255, 255, 255), width=3)
        text = str(number)
        box = draw.textbbox((0, 0), text, font=label_font)
        draw.text((x-(box[2]-box[0])/2, y-(box[3]-box[1])/2-2), text, font=label_font, fill=(0, 15, 24, 255))
    image.convert("RGB").save(FIGURE, quality=94)
    return True


def make_settings_figure():
    if not SETTINGS_SCREENSHOT.exists():
        return False
    SETTINGS_FIGURE.parent.mkdir(parents=True, exist_ok=True)
    image = Image.open(SETTINGS_SCREENSHOT).convert("RGBA")
    draw = ImageDraw.Draw(image, "RGBA")
    try:
        label_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 30)
    except OSError:
        label_font = ImageFont.load_default()
    points = [
        (0.12, 0.31), (0.13, 0.51), (0.13, 0.55), (0.29, 0.59),
        (0.49, 0.70), (0.18, 0.88), (0.39, 0.88), (0.67, 0.88),
    ]
    radius = max(20, image.width // 70)
    for number, (xf, yf) in enumerate(points, 1):
        x, y = int(image.width * xf), int(image.height * yf)
        draw.ellipse((x-radius, y-radius, x+radius, y+radius), fill=(0, 196, 220, 235), outline=(255, 255, 255, 255), width=3)
        text = str(number)
        box = draw.textbbox((0, 0), text, font=label_font)
        draw.text((x-(box[2]-box[0])/2, y-(box[3]-box[1])/2-2), text, font=label_font, fill=(0, 15, 24, 255))
    image.convert("RGB").save(SETTINGS_FIGURE, quality=94)
    return True


def make_workspace_diagram(key, title, zones):
    """Create a clean, manual-style schematic when a live screenshot would date quickly."""
    path = WORKSPACE_ASSETS / f"{key}.png"
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 820), (4, 18, 28))
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 34)
        label_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 25)
        small_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 19)
    except OSError:
        title_font = label_font = small_font = ImageFont.load_default()
    draw.rectangle((22, 22, 1578, 798), outline=(0, 196, 220), width=3)
    draw.text((55, 45), title.upper(), font=title_font, fill=(232, 238, 245))
    draw.text((55, 95), "SCHEMATIC CONTROL MAP · LIVE VALUES AND BUTTON AVAILABILITY VARY", font=small_font, fill=(91, 150, 170))
    columns = 2 if len(zones) <= 6 else 3
    rows = (len(zones) + columns - 1) // columns
    x_gap, y_gap = 30, 24
    left, top, right, bottom = 55, 145, 1545, 750
    cell_w = (right - left - x_gap * (columns - 1)) // columns
    cell_h = (bottom - top - y_gap * (rows - 1)) // rows
    for index, (name, detail) in enumerate(zones, 1):
        offset = index - 1
        column, row = offset % columns, offset // columns
        x = left + column * (cell_w + x_gap)
        y = top + row * (cell_h + y_gap)
        draw.rounded_rectangle((x, y, x + cell_w, y + cell_h), radius=14, fill=(10, 38, 53), outline=(28, 96, 118), width=3)
        draw.ellipse((x + 18, y + 18, x + 66, y + 66), fill=(0, 196, 220), outline=(232, 238, 245), width=2)
        number = str(index)
        box = draw.textbbox((0, 0), number, font=label_font)
        draw.text((x + 42 - (box[2]-box[0])/2, y + 42 - (box[3]-box[1])/2 - 2), number, font=label_font, fill=(4, 18, 28))
        draw.text((x + 82, y + 18), name.upper(), font=label_font, fill=(232, 238, 245))
        draw.multiline_text((x + 24, y + 82), detail, font=small_font, fill=(145, 173, 188), spacing=7)
    image.save(path, quality=94)
    return path


def add_workspace_diagram(doc, key, title, zones):
    path = make_workspace_diagram(key, title, zones)
    doc.add_picture(str(path), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(caption.add_run(f"Control diagram · {title}."), size=9, color=MUTED)


def add_dashboard_legend(doc):
    rows = [
        ("1", "System status", "Overall readiness and current findings."),
        ("2", "Focused probe", "Selects which probe supplies sector, resource, navigation, and fleet context."),
        ("3", "Primary navigation", "Opens Mission Control, Fleet, Galaxy Map, Navigation, Resources, Missions, Production, Safety, Logbook, and Settings."),
        ("4", "Live sector view", "Star-centered schematic with planets, asteroids, probes, relays, and grouped Manny activity."),
        ("5", "Fleet status", "Fleet totals and operational/active/critical posture."),
        ("6", "Resource summary", "Focused-probe fuel percentage and accessible resource quantities."),
        ("7", "Safety overview", "Highest-priority safety state and plain-language reason."),
        ("8", "Active alerts", "Warnings and critical conditions; open the full Safety view for details."),
        ("9", "Active missions", "A compact summary; click for the full mission list."),
        ("10", "Production queue", "Current crafting, assembly, and mining work; click for full details."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [600, 2150, 6610])
    for i, label in enumerate(("No.", "Control", "Purpose")):
        shade(table.rows[0].cells[i], "E8EEF5")
        font(table.rows[0].cells[i].paragraphs[0].add_run(label), bold=True, color=BLUE)
    for number, control, purpose in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (number, control, purpose)):
            font(cell.paragraphs[0].add_run(text), size=9.5)
        set_table_geometry(table, [600, 2150, 6610])


def add_settings_legend(doc):
    rows = [
        ("1", "Account & API credential", "Secure API-key storage, connection testing, removal, and first-launch walkthrough."),
        ("2", "Automation execution", "Explains the current command mode and its safety boundaries."),
        ("3", "Execution mode", "Observe Only, Require Approval, or Automatic."),
        ("4", "Command allowlist", "Limits which crafting, mining, and travel command families may be sent."),
        ("5", "Proposed command queue", "Shows actionable work or the reason nothing can currently run."),
        ("6", "Automation target", "The object or reserve Skunkworks should maintain."),
        ("7", "Desired quantity", "How many completed objects or how much reserve should exist."),
        ("8", "Priority", "Importance from 1–10; 1 is highest and equal values share priority."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [600, 2450, 6310])
    for i, label in enumerate(("No.", "Settings area", "Purpose")):
        shade(table.rows[0].cells[i], "E8EEF5")
        font(table.rows[0].cells[i].paragraphs[0].add_run(label), bold=True, color=BLUE)
    for number, area, purpose in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (number, area, purpose)):
            font(cell.paragraphs[0].add_run(text), size=9.3)
        set_table_geometry(table, [600, 2450, 6310])


def add_warranty_redemption_form(doc):
    """A printable comic insert that reads like an actual service claim form."""

    doc.add_page_break()
    banner = doc.add_table(rows=1, cols=1)
    set_table_geometry(banner, [9360])
    shade(banner.cell(0, 0), "061722")
    p = banner.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("MANNY LIMITED INTERSTELLAR WARRANTY"), size=18, bold=True, color=CYAN, name="Aptos Display")
    p = banner.cell(0, 0).add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("SERVICE JOB FLYER / REDEMPTION FORM  MW-3000-ECE"), size=10, bold=True, color=LIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run("Submit to the nearest imaginary Skunkworks Warranty Department. Processing: three to five business orbits, excluding relativistic delays, black holes, and lunch."), size=8.6)

    layout = doc.add_table(rows=1, cols=2)
    layout.style = "Table Grid"
    set_table_geometry(layout, [5420, 3940])
    left, right = layout.rows[0].cells
    shade(left, "F7F9FB")
    shade(right, "E8EEF5")
    p = left.paragraphs[0]
    font(p.add_run("CLAIMANT AND UNIT IDENTIFICATION"), size=10, bold=True, color=BLUE)
    for label in (
        "Operator / probe owner: __________________________________",
        "Probe name: ______________________________________________",
        "Manny designation: ________________________________________",
        "Serial / task ID: __________________________________________",
        "ECE mined since last imaginary oil change: _________________",
        "Current sector (FCC): ________ / ________ / ________",
    ):
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(label), size=8.5)
    p = left.add_paragraph()
    font(p.add_run("SERVICE REQUEST - CHECK ALL THAT APPLY"), size=10, bold=True, color=BLUE)
    for label in (
        "[ ] Sympathetic nod after avoidable fuel-floor warning",
        "[ ] 3,000-ECE oil change (oil port location unknown)",
        "[ ] Cargo-detachment logistics counseling",
        "[ ] Sensor head points at things judgmentally",
        "[ ] Arm makes noise best described as 'expensive'",
        "[ ] Return-trip definition requires arbitration",
        "[ ] Other: ___________________________________________",
    ):
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        font(p.add_run(label), size=8.1)
    if MANNY_DIAGRAM.exists():
        paragraph = right.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(MANNY_DIAGRAM), width=Inches(1.95))
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("MANNY FIELD-INSPECTION DIAGRAM"), size=9, bold=True, color=BLUE)
    for label in (
        "A. Sensor head - sees warning; cannot make operator read it",
        "B. Utility arm - rated for tools, cargo, and expressive shrugging",
        "C. Service bay - oil port still not found",
        "D. Mobility assembly - warranty void where gravity disagrees",
    ):
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        font(p.add_run(label), size=7.5)

    add_heading(doc, "Description of alleged defect", 2)
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        font(p.add_run("________________________________________________________________________________"), size=8)

    add_heading(doc, "Warranty adjuster checklist", 2)
    checklist = doc.add_table(rows=1, cols=1)
    set_table_geometry(checklist, [9360])
    shade(checklist.cell(0, 0), "F7F9FB")
    checklist.cell(0, 0).paragraphs[0].clear()
    for text in (
        "[ ] Fuel floor reviewed—or claimant has prepared a convincing story.",
        "[ ] Emergency Stop treated as a control, not decorative lighting.",
        "[ ] Five-container detachment recorded as an unscheduled logistics exercise.",
        "[ ] No black-hole travel affecting the probe or the definition of 'return trip.'",
    ):
        p = checklist.cell(0, 0).add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), size=7.8)
    add_note(doc, "Actual terms", "Skunkworks is an independent community tool. The game server remains authoritative, and no joke on this page overrides safety warnings, the software license, or game rules.\nOperator signature: __________________________    Date / local orbit: __________________", "warning")


def add_contents_page(doc, entries):
    add_heading(doc, "Contents")
    add_body(doc, "The manual follows the application from left to right. Page ranges include every subsection and screenshot belonging to that top-level tab.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1200, 6360, 1800])
    for index, label in enumerate(("Section", "Application area", "Pages")):
        shade(table.rows[0].cells[index], "061722")
        font(table.rows[0].cells[index].paragraphs[0].add_run(label), bold=True, color=LIGHT)
    for number, title, pages in entries:
        cells = table.add_row().cells
        font(cells[0].paragraphs[0].add_run(number), bold=True, color=BLUE)
        font(cells[1].paragraphs[0].add_run(title), bold=True if number != "" else False)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(cells[2].paragraphs[0].add_run(pages), bold=True, color=CYAN)
        set_table_geometry(table, [1200, 6360, 1800])


def build():
    has_figure = make_dashboard_figure()
    has_settings_figure = make_settings_figure()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1.0)
    section.left_margin = section.right_margin = Inches(1.0)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, CYAN),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("SKUNKWORKS  |  OPERATOR MANUAL  |  v0.3"), size=8, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Autonomous Exploration & Fleet Operations  |  Updated 2026-08-22"), size=8, color=MUTED)

    # Editorial-cover opening, adapted to the Skunkworks visual language.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("S K U N K W O R K S"), size=28, bold=True, color=NAVY, name="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("OPERATOR MANUAL"), size=18, bold=True, color=CYAN, name="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    font(p.add_run("Mission Control, fleet automation, navigation, resources, and safety"), size=12, color=MUTED)
    add_screen_figure(doc, "mission-control", "Mission Control: live fleet, sector, resources, safety, alerts, and production at a glance.")
    add_note(doc, "Edition", "Version 0.3 is the illustrated living manual for the current development build. Screens and capabilities will be revised with the application.")
    doc.add_page_break()

    add_contents_page(doc, [
        ("1", "Getting Started", "3"),
        ("2", "Mission Control", "4–5"),
        ("3", "Fleet", "6"),
        ("4", "Galaxy Map", "7–8"),
        ("5", "Navigation", "9–11"),
        ("6", "Resources", "12–13"),
        ("7", "Missions", "14"),
        ("8", "Production", "15"),
        ("9", "Safety", "16"),
        ("10", "Communications", "17"),
        ("11", "Manual Control", "18–22"),
        ("12", "Settings", "23–28"),
        ("13", "Troubleshooting and Glossary", "29"),
        ("", "Manny Warranty Redemption Form", "30"),
    ])
    doc.add_page_break()

    add_heading(doc, "1. Getting Started")
    add_body(doc, "Skunkworks is a command-and-awareness layer for Von Neumann Game. It combines live account data, locally retained discoveries, planning, safety review, and optional automation. The focused probe determines the operational context shown across most screens.")
    add_note(doc, "Command authority", "Observe Only never sends game orders. Approval requires confirmation. Automatic can send allowlisted orders when live execution is enabled. The emergency stop overrides every mode.", "warning")
    add_heading(doc, "First launch", 2)
    add_steps(doc, [
        "Open Settings and enter the game API key. Skunkworks stores it in the operating-system credential vault, not in its settings file or logs.",
        "Use Test Connection and confirm that the network status becomes Connected.",
        "Select each probe from the Focused Probe menu and verify its name, model, status, sector, and fuel.",
        "Keep automation in Observe Only while reviewing the dashboard, maps, resources, and proposed command queue.",
        "Choose an execution mode and command allowlist only after the proposed work matches your intent.",
    ])
    add_heading(doc, "Status language", 2)
    add_bullets(doc, [
        "Nominal / Ready — no warning or critical condition was found. Informational notices may still reduce the readiness score slightly.",
        "Degraded — operation can continue, but at least one warning-level resource, data-quality, or safety condition needs attention.",
        "Warning — review is recommended before issuing a related command.",
        "Critical — a serious hazard or stop condition is present.",
        "Unknown — the API or local scan history does not contain enough information to make a reliable claim.",
    ])
    add_note(doc, "Notices versus health", "A notice such as every Manny currently being occupied remains visible, but it does not by itself mark the entire system Degraded. Open Safety or the finding details to see what affected the readiness score.")

    doc.add_page_break()
    add_heading(doc, "2. Mission Control Dashboard")
    add_body(doc, "The dashboard is an overview, not a replacement for the detailed workspaces. Its panels summarize the focused probe and fleet; clicking Active Missions, Production Queue, or Alerts opens the corresponding detailed view.")
    add_screen_figure(doc, "mission-control", "Figure 2-1. Mission Control dashboard and primary controls.", width=6.35)
    add_dashboard_legend(doc)

    add_heading(doc, "3. Focused Probe and Fleet")
    add_heading(doc, "Changing operational context", 2)
    add_steps(doc, [
        "Open Focused Probe in the top-right header.",
        "Choose a probe. Skunkworks loads that probe’s authoritative snapshot before replacing the dashboard context.",
        "Confirm the displayed sector, fuel, resources, and sector objects changed to the selected probe.",
        "Use Refresh if the game was changed in another window or the snapshot is stale.",
    ])
    add_note(doc, "Avoid stale assumptions", "A probe name changing immediately does not prove its sector snapshot has loaded. Wait for refresh completion and verify the FCC coordinates.", "warning")
    add_heading(doc, "Fleet workspace", 2)
    add_bullets(doc, [
        "Review every probe, including idle and unreachable probes.",
        "Rename a probe when the game API permits it.",
        "Assign an operational role such as hub, miner, transport, tanker reserve, explorer, or builder support.",
        "Treat roles as Skunkworks planning instructions; they do not change the physical probe model.",
    ])
    add_screen_figure(doc, "fleet", "Figure 3-1. Fleet identity, Manny auto-naming, and selectable probe cards.")
    add_note(doc, "Manny service interval", "Change the oil in your Manny every 3,000 ECE mined, or the completely imaginary warranty department may look sternly in your direction.")

    add_heading(doc, "4. Sector and Galaxy Maps")
    add_heading(doc, "Live sector view", 2)
    add_body(doc, "The sector view is star-centered. One orbital line is shown for each known planet, and planets are placed on their corresponding line. Asteroids, relays, probes, and containers occupy schematic positions because the game API does not expose precise real-time orbital coordinates.")
    add_bullets(doc, [
        "Mannys mining a known asteroid are grouped at that asteroid.",
        "Large Manny populations are summarized with counts to prevent the map from overflowing.",
        "Unknown or incomplete scans must remain visibly identified as uncertain.",
    ])
    add_screen_figure(doc, "mission-control", "Figure 4-1. The current sector schematic anchors probes, Mannys, asteroids, planets, and SCUT infrastructure.", crop=(0.18, 0.17, 0.99, 0.69))
    add_heading(doc, "Galaxy map", 2)
    add_body(doc, "The galaxy map uses FCC X/Y/Z coordinates and initially centers its camera on the focused probe's current sector. Left-drag to rotate, right- or middle-drag to pan, and use the wheel to zoom. The pan arrow buttons provide precise movement, while Center Probe restores the focused-sector view. Lines represent verified neighboring-sector relationships. Select a sector dot to open its detail panel.")
    add_note(doc, "Local knowledge", "Detailed sector history is retained by Skunkworks after scans. A sector cannot display information the game API never exposed or Skunkworks never observed.")
    add_screen_figure(doc, "galaxy", "Figure 4-2. Rotatable FCC galaxy map with discovery, resource, hazard, route, and SCUT filters.")

    doc.add_page_break()
    add_heading(doc, "5. Navigation and Scanning")
    add_heading(doc, "Manual travel", 2)
    add_steps(doc, [
        "Select the probe that will travel.",
        "Enter destination FCC coordinates or choose a discovered sector.",
        "Preview the route and review deuterium, distance, cargo-detachment, and destination-refueling warnings.",
        "Confirm the command only after the displayed return or recovery plan is acceptable.",
    ])
    add_screen_figure(doc, "navigation-travel", "Figure 5-1. Manual destination entry and route review remain one-time commands.")
    add_heading(doc, "Autonomous transport routes", 2)
    add_body(doc, "Transport routes define separate loading and unloading sectors, resource type, loading target, unloading threshold, return behavior, and a protected deuterium floor. Tanker unloading prefers a designated in-sector reserve tanker before filling a general probe when that policy is configured.")
    add_screen_figure(doc, "settings-transport", "Figure 5-2. Recurring transport roles define pickup, delivery, return, cargo thresholds, and protected fuel.")
    add_heading(doc, "Scanning", 2)
    add_bullets(doc, [
        "Scan one adjacent sector when you need a specific observation.",
        "Scan All 12 Neighboring Sectors mirrors the game’s neighborhood scan and saves the available results.",
        "An exploration-role probe can automatically request the neighborhood scan after arriving in a new sector when exploration automation is enabled.",
        "SCUT coverage indicates communication reach; it does not guarantee a detailed scan already exists.",
    ])
    add_screen_figure(doc, "navigation-scan", "Figure 5-3. Neighbor-sector scanning, knowledge state, SCUT coverage, and travel shortcuts.")

    add_heading(doc, "6. Resources and Inventory")
    add_body(doc, "Resources are grouped by where they exist: focused-probe storage, drifting containers, placed containers when exposed by the API, and remaining asteroid contents from the latest detailed scan. Equipment and constructed items must remain visible alongside bulk resources.")
    add_body(doc, "The Inventory & Containers page also provides manual controls for stock moves, resource or item jettison, container deployment to space, asteroids, or planets, detached-object recovery, same-sector deuterium transfer, and Manny reassignment. Live jettison, deployment, recovery, fuel transfer, and reassignment orders require confirmation. A container with space reserved for active crafting output cannot be detached, dropped, transferred, or emptied until that craft completes or is cancelled; the game enforces this reservation and Skunkworks preserves its explanatory error.")
    add_note(doc, "Container operations", "A Manny can recover drifting or detected asteroid-hidden containers. Whole containers can transfer directly to another owned probe in the same sector. Individual items still use jettison and salvage. Planet deployment consumes an Atmospheric Drop Kit, and transferring a busy Manny cancels its active task.")
    add_body(doc, "A manual mining order can deliver to the probe or a selected detached container. Probe delivery follows saved routing rules: a container assigned to the mined resource is preferred before unassigned storage. Automatic remote mining likewise prefers a compatible resource-designated detached container, then an unassigned container with free capacity.")
    add_bullets(doc, [
        "Probe deuterium is displayed as a percentage of capacity.",
        "Bulk metals, ice, carbon compounds, and other resources use ECE quantities.",
        "Inventory entries identify their containing probe or container whenever the API provides that relationship.",
        "Container policies may reserve a container for one resource type or leave it open to any resource.",
        "Transfer and rename commands remain subject to API capability and safety review.",
    ])
    add_note(doc, "Coverage limitation", "Detached-container contents and planet-dropped container details may be hidden by current observation endpoints. Skunkworks labels unavailable data instead of inventing quantities.")
    add_screen_figure(doc, "manual-cargo", "Figure 6-1. Stock movement, jettison, and handoff controls validate the Manny, source, destination, quantity, and reservations.")
    add_screen_figure(doc, "manual-container", "Figure 6-2. Container deployment, recovery, naming, content routing, and craft-reservation reassignment.")

    doc.add_page_break()
    add_heading(doc, "7. Manual Control")
    add_body(doc, "Manual Control holds one-time operator commands. These actions do not become persistent automation goals, but they still use live API validation, confirmation, inventory reservations, Manny availability, and the current focused-probe context.")
    add_heading(doc, "Production and assembly", 2)
    add_body(doc, "Queue one exposed recipe or assemble a probe from two distinct empty containers, the selected assembly Manny, and every required crafted component. The reference panel expands ordinary recipes into total raw inputs while the probe assembly reference identifies physical components that the game consumes.")
    add_screen_figure(doc, "manual-build", "Figure 7-1. One-time crafting, probe assembly, assembly requirements, and the live crafting reference.")
    add_heading(doc, "Manny field operations", 2)
    add_body(doc, "Field Operations covers manual repair, upgrades, mining, same-sector deuterium transfer, and Manny reassignment. Busy Manny transfer cancels the Manny's current task; moving probes are not eligible until both endpoints arrive.")
    add_screen_figure(doc, "manual-field", "Figure 7-2. Manual repair, upgrade, mining, fuel transfer, and Manny transfer controls.")
    add_heading(doc, "Cargo, infrastructure, and networks", 2)
    add_body(doc, "Infrastructure controls activate and equip SCUT relays, inspect and bookmark objects, refuel at a visible station, and discard stranded cargo only after review. Blueprint sharing requires both owners to participate in the selected active SCUT network.")
    add_screen_figure(doc, "manual-network", "Figure 7-3. Relay activation, beacon installation, inspection, bookmarking, station refueling, and emergency cargo re-dock.")
    add_heading(doc, "Motorized asteroid operations", 2)
    add_body(doc, "Motorization, refueling, launch, and Anatiform sculpting are direct high-impact API commands. Installation consumes the listed components immediately. Asteroid launch cannot be cancelled through the current API and advances one sector per 24 hours toward a terminal outcome.")
    add_screen_figure(doc, "manual-asteroid", "Figure 7-4. Asteroid propulsion installation, refueling, launch, and sculpting controls expose capability locks before review.")

    doc.add_page_break()
    add_heading(doc, "Missions workspace", 2)
    add_body(doc, "Missions lists live game objectives and progress. Open an entry for its available API detail; Skunkworks does not invent hidden quest steps.")
    add_workspace_diagram(doc, "missions", "Missions workspace", [
        ("Mission list", "All active or known objectives for the focused account context."),
        ("Mission state", "Status, progress, and timestamps supplied by the game."),
        ("Detail panel", "Selected mission description and currently exposed requirements."),
        ("Refresh context", "Reload after completing game actions in another window."),
    ])

    doc.add_page_break()
    add_heading(doc, "Production workspace", 2)
    add_body(doc, "Production shows every Manny and printer, including idle workers, with progress, local completion time, countdown, target, and the concise automation reason for work started by Skunkworks. Manual Build Order starts a selected recipe without making it a persistent desired-state target, while still respecting higher-priority allocations.")
    add_screen_figure(doc, "production", "Figure 7-5. Production cards show task, progress, local completion time, countdown, target, delivery behavior, and automation reason.")

    doc.add_page_break()
    add_heading(doc, "8. Settings and Automation")
    add_body(doc, "Settings begins with account access, followed by audio, automation authority, desired-state targets, probe roles, safety floors, and help links. Automation policies, targets, priorities, reserves, and safety floors are stored separately for each focused probe. Owned probe roles are fleet-wide and can only be changed while the main/default probe is focused. Read the section labels before changing values; quantity and priority controls answer different questions.")
    add_body(doc, "Check for Updates opens the official latest-release channel. Download the signed package for the current operating system; the running application never silently replaces itself or bypasses platform security checks. Open Diagnostic Logs opens the per-user support folder containing rotating error logs that can be attached to a bug report. API keys, authorization tokens, and secrets are redacted automatically.")
    add_note(doc, "Updating without losing data", "Stop Skunkworks and use Settings → Back Up Database before updating. Replace only the old application with the new package. Do not delete the platform Skunkworks user-data folder or change SKUNKWORKS_HOME. The database, settings, discovered sectors, roles, routes, and reports are stored separately from the application and should reappear automatically after launch.", "warning")
    add_body(doc, "API compatibility is checked at startup and every six hours. If the server advertises an unreviewed API version, Skunkworks keeps the last valid snapshot visible, labels it stale, and pauses live commands until compatibility has been reviewed. This version check does not consume the probe request budget.")
    add_screen_figure(doc, "settings-policy", "Figure 8-1. Credentials, audio, execution mode, live-order permission, allowlist, cycle limit, and queue state.")
    add_settings_legend(doc)
    add_note(doc, "More settings below", "Scroll to configure probe roles, live target status, resource and safety floors, and the Help & Documentation links.")

    add_heading(doc, "Automation and priorities", 2)
    add_body(doc, "Automation configuration is probe-specific. The mode, live-order permission, command allowlist, maximum orders per cycle, targets, priorities, reserves, and safety floors shown in Settings belong to the focused probe. Switching probes loads that probe's saved configuration and queue; a command prepared for one probe cannot be dispatched under another probe's policy. Probe-role assignments are the fleet-wide exception and are editable only from the main/default probe.")
    add_body(doc, "Automation targets describe a desired state. Priority uses a 1–10 scale: 1 is highest; equal numbers receive equal priority. Numeric selectors can be changed with their minus and plus buttons or by clicking the displayed number, typing a replacement, and pressing Enter. Bounds still apply. The planner compares targets with existing inventory, active work, available Mannys, recipes, resource needs, fuel floors, and safety policy.")
    add_note(doc, "Reading command rows", "Each command row now names its actual output after the command type—for example, MANNY CRAFT · ADDITIONAL CONTAINER or MANNY CRAFT · SCUT RELAY. The P-number is the current saved priority of that specific output target. Several Manny Craft rows with different priorities are distinct targets, not historical copies of one target.")
    add_note(doc, "Allocation ledger", "Active crafting outputs count toward the goal that requested them, preventing duplicate work. During each planning cycle, higher-priority operations claim their required stored resources and component items first. Lower-priority commands cannot spend those claims; equal-priority goals are evaluated in stable target order and refreshed against live inventory before execution.")
    add_screen_figure(doc, "settings-targets", "Figure 8-2. Desired fleet/production quantities and their independent global priorities.")
    add_screen_figure(doc, "settings-floors", "Figure 8-3. Mining-order size, safe travel segment, resource reserves, fuel, free-capacity, and repair floors.")
    add_screen_figure(doc, "settings-status", "Figure 8-4. Live target status compares stored and active work with every configured goal.")
    add_heading(doc, "Reading the planner", 2)
    add_body(doc, "When no command is ready, Complete Planner Status lists every goal in priority order and names the blockers. Missing resources, fabricator availability, fleet-role assignment, fuel, and idle-Manny shortages remain visible instead of being silently discarded.")
    add_screen_figure(doc, "settings-planner", "Figure 8-5. Planner explanations show next-unit requirements, active allocations, uncovered resources, and blockers.")
    add_heading(doc, "Probe roles and logistics", 2)
    add_body(doc, "The default probe assigns fleet-wide roles. Role-specific settings then appear while the corresponding focused probe is selected. Reserve tankers monitor one downstream probe and preserve a protected source reserve; transport and tanker roles define recurring round trips.")
    add_screen_figure(doc, "settings-roles", "Figure 8-6. Default-probe role assignment for the owned fleet.")
    add_screen_figure(doc, "settings-reserve", "Figure 8-7. Reserve tanker refill-chain settings select the monitored consumer and protected source reserve.")
    add_heading(doc, "Recommended commissioning sequence", 2)
    add_steps(doc, [
        "Set target quantities for probes, tankers, Mannys, containers, SCUT relays, and transit beacons.",
        "Assign priorities from 1 to 10. Use equal values when outcomes are equally important.",
        "Set resource reserves and minimum fuel/free-capacity floors.",
        "Choose Observe Only and run one evaluation cycle.",
        "Inspect the proposed command queue and its explanation of why each command exists or why no action is available.",
        "Enable only the command types you are willing to send, then choose Approval or Automatic.",
    ])
    add_note(doc, "Mining quantities", "Maximum per Manny Mining Order is stored separately for each probe and accepts 0.05–0.55 ECE in 0.05 increments. A Manny carries 0.05 ECE per trip, so a larger order remains one continuous multi-trip campaign. Lower the maximum to return Mannys to the available pool sooner and let Skunkworks reconsider other work more frequently; 0.55 ECE preserves the longest, least-interrupted campaign behavior. The planner still tracks the full uncovered requirement and sends another capped order later when it remains necessary.")
    add_heading(doc, "Max orders per cycle", 2)
    add_body(doc, "This limit is the maximum number of separate live game orders Skunkworks may send during one 60-second automatic cycle. It is a rate and blast-radius control, not a target quantity.")

    add_heading(doc, "9. Safety Controls")
    add_bullets(doc, [
        "Travel-distance risk warns when a proposed journey may destroy a probe.",
        "Cargo-detachment risk increases when carrying five or more containers.",
        "Fuel-floor checks consider outbound travel, return travel, and whether either endpoint has a usable deuterium source.",
        "Resource depletion warns before a mining base exhausts its local source and reminds the operator that at most five wandering asteroids may spawn per system.",
        "Safety profiles warn and request acknowledgement; they do not permanently forbid a user-chosen risk unless an explicit stop policy is selected.",
        "Emergency Stop prevents all automation commands until cleared.",
    ])
    add_note(doc, "Observed rules", "Some hazards are observed game behavior rather than formally documented API guarantees. Skunkworks must identify the evidence level and keep thresholds configurable where uncertainty remains.", "warning")

    add_heading(doc, "10. Audio, Logbook, and Daily Use")
    add_heading(doc, "Audio", 2)
    add_bullets(doc, [
        "Background Music enables the cinematic soundtrack and stores its volume separately.",
        "Interface Effects controls clicks, confirmations, warnings, loading, and discovery cues.",
        "Hover Sounds adds subtle feedback to supported navigation, selector, and map targets. It is disabled by default.",
        "Use Test Effect after changing audio output or volume.",
    ])
    add_heading(doc, "Logbook", 2)
    add_body(doc, "The logbook holds user-authored probe notes and, when enabled, significant Skunkworks work or discovery reports. Routine refreshes should not become log entries. Existing game pages should be listed and remain editable or deletable when the API permits those operations.")
    add_screen_figure(doc, "logbook", "Figure 10-1. Focused-probe logbook pages and optional daily role/discovery reporting.")
    add_heading(doc, "Suggested session closeout", 2)
    add_steps(doc, [
        "Review Safety and clear or acknowledge outstanding findings.",
        "Confirm active missions and production work are consistent with current priorities.",
        "Review autonomous transport routes and protected deuterium floors.",
        "Create a logbook summary for discoveries, completed construction, stranded probes, and next-session objectives.",
        "Use Emergency Stop before closing if live automation should not continue.",
    ])
    add_note(doc, "Completion times", "Active work shows the estimated end as a local calendar date, 24-hour time, timezone abbreviation, and UTC offset. A live HH:MM:SS countdown appears beneath the estimate and updates once per second.")
    add_note(doc, "Mining targets", "Production details show the API's current public planet or asteroid name. Internal object identifiers are used only when the API does not provide a public label.")

    add_heading(doc, "11. Troubleshooting Quick Reference")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 3100, 4060])
    for i, label in enumerate(("Symptom", "Check", "Action")):
        shade(table.rows[0].cells[i], "E8EEF5")
        font(table.rows[0].cells[i].paragraphs[0].add_run(label), bold=True, color=BLUE)
    cases = [
        ("Music works; effects are silent", "Interface Effects, effect volume, and Test Effect", "Restart after an audio-device change; verify the Qt FFmpeg backend decodes the selected effect."),
        ("Selected probe shows another probe’s sector", "Focused Probe name and FCC coordinates after refresh", "Refresh; do not act until the authoritative probe snapshot replaces the previous context."),
        ("No automation command is queued", "Targets, inventory, active work, idle assets, allowlist, and safety state", "Read the queue explanation; lower conflicting priorities or supply the missing prerequisite."),
        ("Status is Ready but readiness is below 100%", "Open Safety and inspect notice-level findings", "No action is required for health state; notices such as all Mannies being busy are advisory."),
        ("Qt cocoa platform plugin missing", "Active virtual environment and PySide6 installation", "Use the project `.venv`; avoid mixing Qt/PySide installations from different Python environments."),
        ("Map lacks sector detail", "Knowledge level and last scan", "Scan the sector or its neighbors; unavailable historic detail cannot be reconstructed automatically."),
    ]
    for case in cases:
        cells = table.add_row().cells
        for cell, text in zip(cells, case):
            font(cell.paragraphs[0].add_run(text), size=9.2)
        set_table_geometry(table, [2200, 3100, 4060])

    add_heading(doc, "Glossary", 1)
    for term, definition in (
        ("ECE", "The game’s resource/capacity quantity unit."),
        ("FCC", "Three-dimensional sector coordinates: X, Y, and Z."),
        ("Manny", "Autonomous construction and mining unit."),
        ("SCUT relay", "Communication relay that can participate in the discovered network."),
        ("SCUT transit beacon", "Equipment installed for SCUT transit functionality."),
        ("Desired state", "The quantities and priorities automation attempts to maintain."),
        ("Focused probe", "The probe whose operational context drives the current workspace."),
    ):
        add_body(doc, f"{term}. {definition}", bold_lead=f"{term}. ")

    add_warranty_redemption_form(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


def build_tabbed():
    """Build the manual in the same hierarchy as the application's tabs."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1.0)
    section.left_margin = section.right_margin = Inches(1.0)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, CYAN), ("Heading 2", 13, 14, 7, BLUE), ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("SKUNKWORKS  |  OPERATOR MANUAL  |  v0.5"), size=8, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Autonomous Exploration & Fleet Operations  |  Updated 2026-08-22"), size=8, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("S K U N K W O R K S"), size=28, bold=True, color=NAVY, name="Aptos Display")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("OPERATOR MANUAL"), size=18, bold=True, color=CYAN, name="Aptos Display")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(28)
    font(p.add_run("Organized by application tab and workspace"), size=12, color=MUTED)
    add_screen_figure(doc, "mission-control", "Mission Control: live fleet, sector, resources, safety, alerts, and production at a glance.")
    add_note(doc, "Edition", "Version 0.5 follows the application's navigation hierarchy. Each top-level tab is one section; its sub-tabs, instructions, and supplied screenshots remain together.")
    doc.add_page_break()
    add_contents_page(doc, [
        ("1", "Getting Started", "3"), ("2", "Mission Control", "4"), ("3", "Fleet", "5"),
        ("4", "Galaxy Map", "6"), ("5", "Navigation", "7–9"), ("6", "Resources", "10"),
        ("7", "Missions", "11"), ("8", "Production", "12"), ("9", "Safety", "13"),
        ("10", "Communications", "14"), ("11", "Manual Control", "15–21"),
        ("12", "Settings", "22–31"), ("13", "Troubleshooting and Glossary", "32"),
        ("", "Manny Warranty Redemption Form", "33"),
    ])

    doc.add_page_break(); add_heading(doc, "1. Getting Started")
    add_body(doc, "Skunkworks is a command-and-awareness layer for Von Neumann Game. The focused probe supplies the operational context shown throughout the application.")
    add_note(doc, "Command authority", "Observe Only never sends orders. Approval requires confirmation. Automatic can send allowlisted orders when live execution is enabled. Emergency Stop overrides every mode.", "warning")
    add_heading(doc, "First launch", 2)
    add_steps(doc, [
        "Open Settings and securely save the game API key.", "Test the connection and confirm Connected appears in the header.",
        "Select each probe and verify its name, role, sector, fuel, and inventory.",
        "Keep automation in Observe Only until the proposed queue matches your intent.",
    ])
    add_heading(doc, "Reading the header", 2)
    add_bullets(doc, ["System Status summarizes readiness.", "Focused Probe changes the context for every probe-specific tab.", "Refresh reloads live account data.", "Stop immediately prevents automated orders."])

    doc.add_page_break(); add_heading(doc, "2. Mission Control")
    add_body(doc, "Mission Control is the at-a-glance dashboard for the focused probe. It combines fleet posture, accessible resources, the current sector, alerts, hull integrity, safety, and active production without replacing the detailed tabs.")
    add_screen_figure(doc, "mission-control", "Figure 2-1. Mission Control dashboard and primary navigation.", width=6.35)
    add_dashboard_legend(doc)

    doc.add_page_break(); add_heading(doc, "3. Fleet")
    add_body(doc, "Fleet contains probe identity and fleet-wide selection. Rename probes, choose the default probe when reachable, configure Manny naming for the focused probe, and select a probe card to change context.")
    add_heading(doc, "Probe and Manny naming", 2)
    add_bullets(doc, ["The {probe} field inserts the focused probe name.", "The {number} field inserts the selected numeric or letter sequence.", "Apply to Existing Mannys performs an immediate bulk rename for the focused probe."])
    add_screen_figure(doc, "fleet", "Figure 3-1. Probe identity, Manny auto-naming, direct Manny rename, and fleet cards.")

    doc.add_page_break(); add_heading(doc, "4. Galaxy Map")
    add_body(doc, "Galaxy Map presents the locally known FCC network. Drag to orbit or pan, use the wheel to zoom, and use the view buttons for precise orientation. Fit Map centers and zooms the camera to contain every sector admitted by the current filters. Verified neighboring-sector lines may disappear during camera motion for responsiveness, then return after the view settles at every zoom distance.")
    add_heading(doc, "Filters and sector details", 2)
    add_bullets(doc, ["Filter by discovery state, resource, hazard, dropped container, focused-probe trail, or SCUT coverage.", "Select a node to read its coordinates, observations, objects, confidence, and last visit.", "Scan/Refresh requests current knowledge when the selected sector is reachable by the API."])
    add_note(doc, "Local knowledge", "Detailed history persists after scans, but Skunkworks never invents information the game API did not expose.")
    add_screen_figure(doc, "galaxy", "Figure 4-1. Rotatable FCC galaxy view, sector inspector, filters, trail, and SCUT coverage.")

    doc.add_page_break(); add_heading(doc, "5. Navigation")
    add_heading(doc, "Manual Travel", 2)
    add_body(doc, "Manual Travel prepares one movement command for the focused probe. Enter an FCC directly or choose Use for Manual Travel beside a neighboring scan, then select the route strategy and preview before anything is sent.")
    add_steps(doc, ["Confirm the focused probe and current FCC in the header.", "Enter the destination X, Y, and Z coordinates.", "Choose the segmented route strategy and select Preview Route.", "Review every segment, deuterium requirement, fuel reserve, cargo-detachment warning, and refueling dependency.", "Use Confirm One-Time Travel Command only for this trip."])
    add_note(doc, "One-time order", "Manual confirmation does not create a repeating route. Movement still requires the command allowlist and live execution authority.")
    add_screen_figure(doc, "navigation-travel", "Figure 5-1. Manual Travel keeps destination, preview, warnings, and confirmation together.", width=6.25)
    doc.add_page_break()
    add_heading(doc, "Automatic Travel", 2)
    add_body(doc, "After a segmented route has been previewed, Save Auto-Travel Destination stores that reviewed destination as a durable planner goal. It is not an unconditional launch button: the planner waits until Automatic mode, live orders, and Move Probe permission are all enabled.")
    add_bullets(doc, ["The saved goal remains associated with the focused probe.", "The planner sends one safe segment at a time and rechecks the live probe state after refresh.", "Fuel and safety floors continue to apply before every departure.", "Recurring pickup-and-delivery loops are configured separately under Settings → Probe Role Settings → Transport or Deuterium Tanker."])
    add_screen_figure(doc, "navigation-travel", "Figure 5-2. The same reviewed route can be saved as an automatic destination from Manual Travel.", width=6.25)
    doc.add_page_break()
    add_heading(doc, "Sector Scanning", 2)
    add_body(doc, "Sector Scanning lists the twelve valid FCC neighbors around the focused probe. Each row identifies its knowledge level, estimated contents, and SCUT coverage.")
    add_bullets(doc, ["Scan requests current data for one neighbor.", "Scan All 12 Neighboring Sectors refreshes the complete neighborhood.", "Details opens the retained observation for that FCC.", "Use for Manual Travel copies the selected FCC into the travel form.", "Detailed and Neighbor Scan labels distinguish authoritative observations from lower-detail estimates."])
    add_screen_figure(doc, "navigation-scan", "Figure 5-3. Sector Scanning actions and knowledge states.", width=6.25)

    doc.add_page_break(); add_heading(doc, "6. Resources")
    add_body(doc, "Resources groups probe storage, containers, drifting objects, and the latest observed asteroid reserves. Quantities retain their source location so users can distinguish accessible cargo from remote or detached stock.")
    add_heading(doc, "Reading inventory", 2)
    add_bullets(doc, ["Probe deuterium is displayed as a percentage of tank capacity.", "Bulk resources use ECE quantities and retain their containing probe or container.", "Asteroid quantities are the latest recorded scan, not a prediction of current stock.", "Unknown or unavailable container contents remain explicitly labeled."])
    add_note(doc, "Related controls", "Resource movement, jettison, container routing, deployment, and recovery are operator commands under Manual Control → Cargo and Transfers. They are documented there with their actual screenshots.")
    add_screen_figure(doc, "resources", "Figure 6-1. Resource ledger groups probe storage, containers, detached cargo, and observed natural reserves.", width=6.2)

    doc.add_page_break(); add_heading(doc, "7. Missions")
    add_body(doc, "Missions lists objectives and progress exposed by the game. Select an entry to review its description, current state, progress, and any requirements supplied by the API. Refresh after completing mission work elsewhere so the displayed state is authoritative.")
    add_bullets(doc, ["The mission list is account context, while probe-specific requirements still depend on the focused probe.", "Skunkworks does not infer hidden objectives or undisclosed completion conditions."])
    add_screen_figure(doc, "missions", "Figure 7-1. Missions lists objectives and progress exposed by the game API.", width=6.2)

    doc.add_page_break(); add_heading(doc, "8. Production")
    add_body(doc, "Production displays every Manny and printer, including idle workers. Active cards show the task, target, progress, local completion time, countdown, delivery behavior, and the reason automation dispatched the order.")
    add_bullets(doc, ["Sort Mannys by name, status, or remaining time.", "Idle Ready means the Manny can receive a valid order.", "Recall cancels the current task subject to game rules.", "Mining cards show their target object, resource, amount, and delivery commitment."])
    add_screen_figure(doc, "production", "Figure 8-1. Live Manny production, crafting, mining, countdowns, and recall controls.")

    doc.add_page_break(); add_heading(doc, "9. Safety")
    add_body(doc, "Safety centralizes operational findings. Readiness distinguishes advisory notices from warning and critical conditions; a busy fleet can remain healthy even when no Manny is currently available.")
    add_bullets(doc, ["Travel risk considers distance, fuel, return options, and destination refueling.", "Cargo warnings identify detachment risk at five or more containers.", "Resource depletion warns before local mining sources are exhausted.", "Emergency Stop prevents automation until explicitly cleared."])
    add_note(doc, "Evidence", "Some hazards are observed game behavior rather than formal API guarantees. The application identifies uncertainty and keeps configurable thresholds where appropriate.", "warning")
    add_screen_figure(doc, "safety", "Figure 9-1. Safety groups current findings, severity, and available recovery actions.", width=6.2)

    doc.add_page_break(); add_heading(doc, "10. Communications")
    add_heading(doc, "Messaging", 2)
    add_body(doc, "Messaging contains account communications and unread state exposed by the game API. Messages remain separate from operational alerts and safety findings.")
    add_heading(doc, "Logbook", 2)
    add_body(doc, "Logbook pages belong to the focused probe. Create, edit, and delete operator notes independently of automatic reporting. Routine refreshes are not logged as events.")
    add_heading(doc, "Optional Automatic Reports", 3)
    add_body(doc, "Automatic daily role reports and major-discovery reports are an opt-in feature controlled by the Auto-Log Daily Role Reports and Major Discoveries checkbox in the Logbook sub-tab. The checkbox is off by default and is stored as an account-wide preference rather than a separate setting for each probe.")
    add_bullets(doc, ["When enabled, the first eligible refresh after 17:00 local time creates one due role-specific report for each owned probe.", "A report summarizes the previous reporting window using retained Skunkworks telemetry and commands accepted by the game; unavailable activity is not inferred.", "Major-discovery pages are also created only while this checkbox is enabled.", "Unchecking the option prevents future automatic reports but does not delete pages already created.", "Manual logbook pages remain available whether automatic reporting is enabled or disabled."])
    add_note(doc, "Privacy", "Automatic reports are written into the game's probe logbooks. Review their contents before sharing screenshots or diagnostic material because reports can contain probe names, sectors, inventory activity, and operational history.", "warning")
    add_screen_figure(doc, "logbook", "Figure 10-1. Messaging/Logbook sub-tabs, page list, editor, and automatic daily reports.", width=4.75)

    add_heading(doc, "11. Manual Control")
    add_body(doc, "Manual Control contains one-time operator orders. Every command remains subject to live validation, confirmation, reservations, and Manny availability.")
    add_note(doc, "Review before sending", "Buttons labeled Review open a confirmation summary. Disabled controls indicate that a required Manny, item, object, blueprint, destination, or API capability is unavailable.")
    doc.add_page_break()
    add_heading(doc, "Production and Assembly", 2)
    add_body(doc, "This sub-tab starts a single recipe or assembles one probe without creating an automation target. Manual Build Order needs a recipe and a compatible idle fabricator. Manual Probe Assembly needs a model, an assembly Manny, two distinct empty containers, and every listed model component.")
    add_bullets(doc, ["Queue Build submits one finished item using the game's exposed recipe.", "The crafting reference expands ordinary recipes into the total raw resources for one finished item.", "Probe Assembly Requirements lists the physical crafted components consumed by each probe model.", "Review Probe Assembly rechecks the selected containers, assembly Manny, and components before confirmation.", "The assembly Manny is installed aboard the completed probe and is no longer available on the source probe."])
    add_screen_figure(doc, "manual-build", "Figure 11-1. Production and Assembly inputs, model requirements, and recipe reference.", width=6.2)
    doc.add_page_break()
    add_heading(doc, "Manny Field Operations", 2)
    add_body(doc, "This sub-tab operates directly on the focused probe and its same-sector neighbors.")
    add_bullets(doc, ["Manual Probe Repair selects an idle Manny and a repair percentage; the order remains unavailable when no Manny can act.", "Manual Probe Upgrade lists only unlocked, unfinished upgrades supported by the current owner and probe.", "Manual Mining Order selects the Manny, mineable object, resource, ECE amount, and delivery destination. The amount is capped by Max per Manny Mining Order.", "Same-Sector Probe Transfers can send deuterium or a Manny to an arrived target probe.", "Transferring a busy Manny cancels that Manny's current task; moving probes are excluded until both probes arrive."])
    add_screen_figure(doc, "manual-field", "Figure 11-2. Repair, upgrade, mining, deuterium transfer, and Manny reassignment.", width=6.2)
    doc.add_page_break()
    add_heading(doc, "Cargo and Transfers", 2)
    add_body(doc, "Inventory and Container Control moves or releases cargo from the focused probe. Every transfer requires an available onboard Manny and confirmation.")
    add_bullets(doc, ["Move Stock selects resource or item, amount, source container, and destination container.", "Manual Jettison creates a recoverable sector object; another same-sector probe can recover it for an item handoff.", "SCUT Relay Deployment uses a stored relay, an idle Manny, an Integrated Circuit for activation, and a SCUT Transit Beacon for transit support.", "Inspect Object records object detail without starting mining; Bookmark Target installs a named reference.", "Waiting Cargo emergency re-dock permanently discards resource cargo but returns recoverable objects to the sector."])
    add_screen_figure(doc, "manual-cargo", "Figure 11-3. Stock movement, jettison, relay deployment, inspection, and waiting cargo.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Container Deployment, Recovery, and Routing", 3)
    add_body(doc, "Deploy one attached container to space, an asteroid, a planet, or another supported destination, then recover a visible drifting object with a Manny. Routing rules control where newly received content is placed.")
    add_bullets(doc, ["Choose Available Manny, Attached Container, and Destination before Review Deployment.", "Choose Recovery Manny and Drifting Object before Review Recovery.", "Rename gives each storage location an operational label.", "Preferred Contents reserves normal routing for one resource type or allows any contents.", "Reassign Craft Reservations moves active output reservations only when all reserved output fits elsewhere; otherwise nothing changes."])
    add_screen_figure(doc, "manual-container", "Figure 11-4. Container deployment, recovery, naming, routing, and reservation reassignment.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Infrastructure and Networks", 2)
    add_body(doc, "Infrastructure controls create and use persistent network objects. The relay workflow is sequential: jettison a stored SCUT Relay, refresh until it appears as inactive, activate it with an Integrated Circuit, then install a stocked SCUT Transit Beacon.")
    add_bullets(doc, ["Each Manny relay step takes five minutes.", "Inspect Object and Bookmark Target work on visible sector objects.", "Station Refuel requires a visible deuterium station.", "Share Improvement Blueprint requires a known blueprint, recipient probe, and shared coverage from the selected active SCUT network.", "Repeating a completed blueprint share is idempotent and does not duplicate the recipient alert."])
    add_screen_figure(doc, "manual-network", "Figure 11-5. Relay workflow, inspection, bookmarks, refueling, waiting cargo, and blueprint sharing.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Motorized Asteroid Operations", 2)
    add_body(doc, "These are direct, high-impact asteroid commands. Controls stay locked until the required owner capability or blueprint is known.")
    add_bullets(doc, ["Install Asteroid Propulsion selects an idle Manny and asteroid; the game immediately consumes the listed engine, steel, and deuterium inputs.", "Refuel Motorized Asteroid selects the motorized asteroid and a valid source.", "Launch selects the asteroid, movement mode, and a directly neighboring FCC. Transfer advances one sector per 24 hours and cannot be cancelled through the current API.", "Sculpt Anatiform Asteroid requires the corresponding blueprint and performs the documented two-day operation.", "Detected Active Trajectories lists movement currently visible in the detailed sector scan."])
    add_note(doc, "Irreversible review", "Read the destination and resource summary carefully before confirming installation or launch.", "warning")
    add_screen_figure(doc, "manual-asteroid", "Figure 11-6. Propulsion installation, refueling, launch, sculpting, locks, and trajectories.", width=6.15)

    doc.add_page_break(); add_heading(doc, "12. Settings")
    add_heading(doc, "General Automation", 2)
    add_body(doc, "General Automation owns the focused probe's account connection, audio, execution authority, desired state, safety floors, planner status, and diagnostics.")
    add_note(doc, "Updating without losing data", "Before updating, stop Skunkworks and use Settings → Back Up Database. Replace only the old application with the new package. Do not delete the platform Skunkworks user-data folder or change SKUNKWORKS_HOME. The database, settings, discovered sectors, roles, routes, and reports are stored separately from the application and should reappear automatically after launch.", "warning")
    add_heading(doc, "Account, Audio, and Execution Policy", 3)
    add_bullets(doc, ["The API key is stored in the operating-system credential vault and is not written into settings or logs.", "Test Connection verifies access; Remove deletes the saved credential.", "Audio separately controls background music, interface effects, hover sounds, and volumes.", "Execution Mode selects Observe Only, Approval, or Automatic.", "Allow Skunkworks to Send Game Orders is the master live-order permission.", "Command Allowlist limits crafting/assembly, mining, travel, deuterium transfer, and repair families.", "Max Orders per 1-Minute Cycle limits the number of commands sent in one automatic cycle."])
    add_screen_figure(doc, "settings-policy", "Figure 12-1. Credentials, audio, execution authority, allowlist, cycle limit, and command queue.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Targets and Priorities", 3)
    add_body(doc, "Fleet Assembly Targets and Production Targets define persistent desired quantities. A target counts stored output plus output already active or allocated.")
    add_bullets(doc, ["Desired Quantity is the completed quantity to maintain, not a one-time batch size.", "Priority 1 is highest; 10 is lowest.", "Higher-priority goals claim current raw resources and required stored components before lower-priority goals.", "Equal-priority goals use stable target order and are reconsidered after each planning cycle.", "Ordinary crafts use the game's direct recipe and reserve only the next unit's raw inputs; probe assembly reserves its required physical components."])
    add_screen_figure(doc, "settings-targets", "Figure 12-2. Persistent fleet and production targets with independent priorities.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Resource and Safety Floors", 3)
    add_bullets(doc, ["Max per Manny Mining Order caps one continuous mining campaign from 0.05–0.55 ECE.", "Safe Segment Length limits how many collision-safe sectors one planned travel segment may span.", "Deuterium, Metals, Ice, and Carbon Compounds define minimum onboard reserves and independent priorities.", "Fuel Floor protects the percentage required before normal work or departure.", "Min Free Capacity keeps storage available for incoming work.", "Auto Repair At / Below and Repair Target define when automatic repair begins and where it stops."])
    add_screen_figure(doc, "settings-floors", "Figure 12-3. Mining size, travel, resource, fuel, capacity, and repair floors.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Live Target Status and Planner", 3)
    add_body(doc, "Live Target Status is the compact desired-state summary. Each row shows current or active quantity, target, remaining work, and priority color. Refresh Diagnostics identifies the slowest refresh stages.")
    add_screen_figure(doc, "settings-status", "Figure 12-4. Live Target Status, remaining work, and refresh diagnostics.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Complete Planner Status", 3)
    add_body(doc, "When no command is ready, the planner still lists every goal. Read each row from the intended output through its next-unit requirements and Waiting For reasons.")
    add_bullets(doc, ["Missing Resources shows the exact uncovered raw quantity after onboard and inbound stock.", "Fabricator Unavailable means the required worker type is occupied or absent.", "No Idle Manny prevents work that specifically requires a Manny.", "Fleet Role Assignment Not Automated identifies a role dependency that must be configured by the operator.", "Higher-priority reservations remain visible so lower-priority work is not mistaken for forgotten work."])
    add_screen_figure(doc, "settings-planner", "Figure 12-5. Planner requirements, allocations, uncovered resources, and named blockers.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Probe Role Settings", 2)
    add_body(doc, "Probe roles are configured together in this sub-tab. The default probe assigns one role to each owned probe. Selecting a role-bearing probe then exposes that role's own controls; roles without settings clearly say that no probe-specific controls are available yet.")
    add_heading(doc, "Role Assignment", 3)
    add_bullets(doc, ["Hub marks the main coordination probe.", "Explorer supports travel and scanning workflows.", "Transport carries a selected resource between configured endpoints.", "Deuterium Tanker carries fuel using an absolute tank target.", "Deuterium Reserve monitors and replenishes a selected downstream probe or tanker."])
    add_screen_figure(doc, "settings-roles", "Figure 12-6. Fleet-wide role assignment from the default probe.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Deuterium Reserve", 3)
    add_body(doc, "Select Monitor and Refill Probe to identify the next link in the reserve chain. Protected Source Reserve is the ECE the reserve tanker must retain before it can forward fuel. Multiple reserve tankers can be chained by selecting another tanker as the monitored consumer.")
    add_screen_figure(doc, "settings-reserve", "Figure 12-7. Reserve-chain consumer and protected source reserve.", width=6.15)
    doc.add_page_break()
    add_heading(doc, "Transport and Deuterium Tanker", 3)
    add_body(doc, "Both recurring roles use Loading Sector, Unloading Sector, Return Point, source/destination probes, repetition, protected deuterium, contingency hops, and an optional verified refuel stop.")
    add_bullets(doc, ["Transport selects a resource, loads until a cargo percentage, and unloads until a percentage remains.", "Deuterium Tanker always carries deuterium and loads the tank to an absolute ECE target.", "Repeat Until Paused makes the saved operation recurring.", "Route Depends on Refueling requires the selected intermediate FCC and minimum source quantity to be verified.", "Review validates the current route; Pause Route stops future departures without deleting the saved plan."])
    add_screen_figure(doc, "settings-transport", "Figure 12-8. General Transport round-trip configuration.", width=6.05)
    doc.add_page_break()
    add_heading(doc, "Deuterium Tanker", 3)
    add_body(doc, "Tanker mode replaces percentage cargo loading with Load Tank To ECE while preserving the protected fuel floor and contingency reserve for the complete loop.")
    add_screen_figure(doc, "settings-tanker", "Figure 12-9. Deuterium Tanker absolute fuel-load configuration.", width=6.05)

    doc.add_page_break(); add_heading(doc, "13. Troubleshooting and Glossary")
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; set_table_geometry(table, [2200, 3100, 4060])
    for i, label in enumerate(("Symptom", "Check", "Action")):
        shade(table.rows[0].cells[i], "E8EEF5"); font(table.rows[0].cells[i].paragraphs[0].add_run(label), bold=True, color=BLUE)
    for case in (
        ("Wrong sector after probe selection", "Focused Probe and FCC after refresh", "Wait for the authoritative snapshot before acting."),
        ("No automation command", "Planner blockers, inventory, allowlist, safety", "Read Complete Planner Status and resolve the named prerequisite."),
        ("Idle Manny receives no work", "Queue, raw resources, fabricator, reservations", "Confirm a command is unblocked and the required worker type is available."),
        ("Refresh is slow", "Refresh Diagnostics slowest stages", "Capture diagnostic logs; galaxy refreshes are expected to be the heaviest."),
        ("Map lacks detail", "Knowledge level and last scan", "Scan the sector or its neighbors."),
    ):
        cells = table.add_row().cells
        for cell, text in zip(cells, case): font(cell.paragraphs[0].add_run(text), size=9.2)
        set_table_geometry(table, [2200, 3100, 4060])
    add_heading(doc, "Glossary", 2)
    for term, definition in (("ECE", "Resource and capacity quantity unit."), ("FCC", "Three-dimensional sector coordinates."), ("Manny", "Autonomous construction and mining unit."), ("SCUT relay", "Communication relay in a discovered network."), ("Desired state", "The quantities and priorities automation maintains."), ("Focused probe", "The probe supplying the current workspace context.")):
        add_body(doc, f"{term}. {definition}", bold_lead=f"{term}. ")
    add_warranty_redemption_form(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_tabbed()
